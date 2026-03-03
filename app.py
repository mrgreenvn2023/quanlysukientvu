import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

st.set_page_config(page_title="TVU Event OS - Template Engine", layout="wide")

# --- 1. KHO DỮ LIỆU MẪU CỦA ĐẠI HỌC TRÀ VINH ---
# Hệ thống không cần AI mạng, tự rút dữ liệu từ kho nội bộ
TVU_TEMPLATES = {
    "Lễ công bố / Kỷ niệm lớn": {
        "kh": [
            {"Hạng mục": "Soạn thảo văn bản", "Nội dung": "Lập tờ trình, xin ý kiến BGH, mời đại biểu", "Phụ trách": "Văn phòng", "Hạn": "Ngày -15"},
            {"Hạng mục": "Thiết kế nhận diện", "Nội dung": "Thiết kế Backdrop, Standee, Thẻ đeo", "Phụ trách": "Tổ Truyền thông", "Hạn": "Ngày -10"},
            {"Hạng mục": "Sản xuất quà tặng", "Nội dung": "In ấn túi tote canvas in logo TVU, chuẩn bị quà VIP", "Phụ trách": "Văn phòng", "Hạn": "Ngày -7"},
            {"Hạng mục": "Thi công sân khấu", "Nội dung": "Dựng khung backdrop, test màn hình LED, âm thanh", "Phụ trách": "Phòng QT-TB", "Hạn": "Ngày -2"},
            {"Hạng mục": "Tổng duyệt kịch bản", "Nội dung": "Chạy thử MC, test nhạc lễ, khớp đội hình lễ tân", "Phụ trách": "Đạo diễn / MC", "Hạn": "Ngày -1"}
        ],
        "kb": [
            {"Giờ": "07:30 - 08:00", "Nội dung": "Đón khách, cài hoa đại biểu, mời dùng Teabreak", "Kỹ thuật": "Nhạc hòa tấu nhẹ nhàng", "Điều phối": "Lễ tân"},
            {"Giờ": "08:00 - 08:10", "Nội dung": "Ổn định tổ chức, văn nghệ chào mừng (2 tiết mục)", "Kỹ thuật": "Nhạc nền sôi động, Ánh sáng full", "Điều phối": "Đoàn Thanh niên"},
            {"Giờ": "08:10 - 08:15", "Nội dung": "Chào cờ, Tuyên bố lý do, Giới thiệu đại biểu", "Kỹ thuật": "Nhạc Quốc ca chuẩn", "Điều phối": "MC"},
            {"Giờ": "08:15 - 08:30", "Nội dung": "Phát biểu khai mạc của Lãnh đạo Trường", "Kỹ thuật": "Micro bục, Slide nền TVU", "Điều phối": "BGH"},
            {"Giờ": "08:30 - 08:45", "Nội dung": "Nghi thức công bố quyết định quan trọng", "Kỹ thuật": "Nhạc hành khúc dồn dập, Pháo hoa điện", "Điều phối": "MC / Kỹ thuật"},
            {"Giờ": "08:45 - 09:00", "Nội dung": "Tặng hoa, trao quà lưu niệm (Túi tote TVU)", "Kỹ thuật": "Nhạc trao thưởng", "Điều phối": "Lễ tân"}
        ],
        "dt": [
            {"Khoản mục": "Sản xuất túi Tote TVU (Quà tặng)", "Số lượng": 500, "Đơn giá": 35000, "Thành tiền": 17500000},
            {"Khoản mục": "Teabreak đón khách VIP", "Số lượng": 100, "Đơn giá": 150000, "Thành tiền": 15000000},
            {"Khoản mục": "In ấn Backdrop & Standee", "Số lượng": 1, "Đơn giá": 5000000, "Thành tiền": 5000000},
            {"Khoản mục": "Thuê âm thanh, ánh sáng chuyên nghiệp", "Số lượng": 1, "Đơn giá": 12000000, "Thành tiền": 12000000},
            {"Khoản mục": "Hoa tươi bục phát biểu & đại biểu", "Số lượng": 10, "Đơn giá": 500000, "Thành tiền": 5000000}
        ]
    },
    "Hội thảo khoa học": {
        "kh": [
            {"Hạng mục": "Mời diễn giả", "Nội dung": "Gửi thư mời các nhà khoa học, chốt chủ đề", "Phụ trách": "Khoa chuyên môn", "Hạn": "Ngày -30"},
            {"Hạng mục": "Thu thập tham luận", "Nội dung": "Nhận bài, biên tập kỷ yếu hội thảo", "Phụ trách": "Phòng NCKH", "Hạn": "Ngày -15"}
        ],
        "kb": [
            {"Giờ": "08:00 - 08:30", "Nội dung": "Check-in, nhận kỷ yếu và túi tài liệu", "Kỹ thuật": "Nhạc không lời", "Điều phối": "Lễ tân"},
            {"Giờ": "08:30 - 09:30", "Nội dung": "Báo cáo tham luận chính (Keynote)", "Kỹ thuật": "Trình chiếu Slide", "Điều phối": "Chủ tọa"}
        ],
        "dt": [
            {"Khoản mục": "In kỷ yếu hội thảo", "Số lượng": 200, "Đơn giá": 80000, "Thành tiền": 16000000},
            {"Khoản mục": "Thù lao báo cáo viên", "Số lượng": 5, "Đơn giá": 2000000, "Thành tiền": 10000000}
        ]
    }
}

# --- 2. HÀM XUẤT VĂN BẢN WORD CHUẨN ---
def export_word(name, df_kh, df_kb, df_dt):
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    l_cell = table.cell(0, 0).paragraphs[0]
    l_cell.add_run("UBND TỈNH TRÀ VINH\nĐẠI HỌC TRÀ VINH\n").bold = True
    l_cell.add_run("Số:      /KH-ĐHTV").italic = True
    r_cell = table.cell(0, 1).paragraphs[0]
    r_cell.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n").bold = True
    r_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"\nKẾ HOẠCH VÀ KỊCH BẢN\n{name.upper()}").bold = True

    for title, df in [("I. KẾ HOẠCH TỔNG THỂ", df_kh), ("II. KỊCH BẢN ĐIỀU HÀNH", df_kb), ("III. DỰ TOÁN KINH PHÍ", df_dt)]:
        doc.add_heading(title, level=1)
        t = doc.add_table(rows=1, cols=len(df.columns))
        t.style = 'Table Grid'
        for i, col in enumerate(df.columns): t.rows[0].cells[i].text = col
        for _, row in df.iterrows():
            row_cells = t.add_row().cells
            for i, val in enumerate(row): row_cells[i].text = str(val)

    doc.add_paragraph("\nNơi nhận:\n- BGH (để b/c);\n- Các đơn vị;\n- Lưu VP.")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# --- 3. GIAO DIỆN CHÍNH ---
st.title("🏛️ TVU EVENT OS (Template Engine)")
st.info("Hệ thống vận hành không cần API. Nhanh chóng, bảo mật và ổn định 100%.")

col1, col2 = st.columns([2, 1])
with col1:
    ev_name = st.text_input("Nhập tên sự kiện cụ thể:", value="Lễ công bố chuyển đổi Đại học Trà Vinh")
with col2:
    loai_su_kien = st.selectbox("Chọn khung kịch bản mẫu:", list(TVU_TEMPLATES.keys()))

if st.button("🚀 NẠP DỮ LIỆU KỊCH BẢN CHUẨN"):
    # Rút dữ liệu từ kho mẫu nội bộ
    st.session_state.kh = pd.DataFrame(TVU_TEMPLATES[loai_su_kien]["kh"])
    st.session_state.kb = pd.DataFrame(TVU_TEMPLATES[loai_su_kien]["kb"])
    st.session_state.dt = pd.DataFrame(TVU_TEMPLATES[loai_su_kien]["dt"])
    st.success("Đã tải xong kịch bản chuẩn! Chuyên gia có thể tinh chỉnh số liệu ngay bên dưới.")

if 'kh' in st.session_state:
    t1, t2, t3, t4 = st.tabs(["📅 KẾ HOẠCH", "🎬 KỊCH BẢN", "💰 DỰ TOÁN", "📄 XUẤT FILE WORD"])
    
    with t1:
        st.write("Chỉnh sửa đầu việc, phân công người phụ trách:")
        st.session_state.kh = st.data_editor(st.session_state.kh, num_rows="dynamic", use_container_width=True)
    with t2:
        st.write("Điều chỉnh đường dây kịch bản (MC, Âm thanh, Ánh sáng):")
        st.session_state.kb = st.data_editor(st.session_state.kb, num_rows="dynamic", use_container_width=True)
    with t3:
        st.write("Điều chỉnh số lượng và đơn giá thực tế:")
        df_dt = st.data_editor(st.session_state.dt, num_rows="dynamic", use_container_width=True)
        try:
            df_dt['Thành tiền'] = pd.to_numeric(df_dt['Số lượng']) * pd.to_numeric(df_dt['Đơn giá'])
            st.session_state.dt = df_dt
            st.metric("TỔNG TIỀN DỰ TOÁN", f"{df_dt['Thành tiền'].sum():,.0f} VNĐ")
        except: pass
    with t4:
        st.write("Hệ thống sẽ tổng hợp 3 bảng trên thành file Word chuẩn thể thức Nghị định 30.")
        if st.button("📦 XUẤT FILE WORD TỔNG HỢP"):
            word_file = export_word(ev_name, st.session_state.kh, st.session_state.kb, st.session_state.dt)
            st.download_button("📥 TẢI FILE WORD CHUẨN", word_file, f"Kich_ban_{ev_name}.docx")
